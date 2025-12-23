
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def _add_page_field(paragraph, field_name: str):
    """Add a Word field (PAGE / NUMPAGES) to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_name
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "1"
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)

def apply_header_footer(doc, header_left: str, header_center: str, header_right: str,
                        version_text: str, effective_date_text: str):
    """
    Apply a clean header/footer similar to regulated lab forms.
    Content kept as placeholders so user can match Excel exactly.
    """
    section = doc.sections[0]
    # HEADER
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].clear()
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(header_center)
    r.bold = True
    r.font.size = Pt(10)

    # optional second line (left/right)
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(header_left)
    r2.font.size = Pt(9)

    # FOOTER
    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{version_text}    |    {effective_date_text}    |    Trang ")
    run.font.size = Pt(8)
    _add_page_field(fp, "PAGE")
    fp.add_run(" / ").font.size = Pt(8)
    _add_page_field(fp, "NUMPAGES")
