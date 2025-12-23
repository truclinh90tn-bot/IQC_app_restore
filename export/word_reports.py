
import io
from dataclasses import dataclass
from typing import Optional, Tuple

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import matplotlib.pyplot as plt

from export.docx_layout import apply_header_footer


@dataclass
class ReportMeta:
    don_vi: str = "{{DON_VI}}"
    phien_ban: str = "Phiên bản: {{PHIEN_BAN}}"
    ngay_hieu_luc: str = "Ngày hiệu lực: {{NGAY_HIEU_LUC}}"
    ten_xet_nghiem: str = ""
    thiet_bi_phuong_phap: str = ""
    lo_qc_han_dung: str = ""
    thang_nam: str = ""


def _safe_str(x) -> str:
    if x is None:
        return ""
    if isinstance(x, float) and (np.isnan(x) or np.isinf(x)):
        return ""
    return str(x)


def build_lj_figure_from_z(z_df: pd.DataFrame,
                           point_df: Optional[pd.DataFrame] = None,
                           title: str = "Levey–Jennings (Z-score)") -> plt.Figure:
    """
    Vẽ Levey–Jennings kiểu giống chart trong app:
    - 3 mức QC là 3 đường
    - đường ngang 0, ±1, ±2, ±3
    - khoanh đỏ các điểm có rule_codes (vi phạm/cảnh báo)
    """
    if z_df is None or z_df.empty:
        raise ValueError("z_df is empty")

    runs = z_df["Ngày/Lần"].astype(str).tolist()
    z_cols = [c for c in z_df.columns if c.startswith("z_Ctrl")]
    z_cols = sorted(z_cols, key=lambda x: int(x.split("Ctrl ")[1]))

    Z = z_df[z_cols].to_numpy(dtype=float)
    n_runs, n_levels = Z.shape

    # build long for violations
    viol = set()
    short_map = {}
    if point_df is not None and not point_df.empty:
        for _, r in point_df.iterrows():
            run = str(r.get("Ngày/Lần", ""))
            ctrl = str(r.get("Control", ""))
            codes = _safe_str(r.get("rule_codes", "")).strip()
            if codes:
                viol.add((run, ctrl))
                # short already computed in app chart (rule_short)
                s = _safe_str(r.get("rule_short", "")).strip()
                if s:
                    short_map[(run, ctrl)] = s

    fig = plt.figure(figsize=(8.2, 4.6), dpi=200)  # ~ 3/4 A4 when inserted
    ax = fig.add_subplot(111)

    x = np.arange(n_runs)
    for lvl in range(n_levels):
        y = Z[:, lvl]
        ax.plot(x, np.clip(y, -3, 3), marker="o", linewidth=1.6, label=f"Ctrl {lvl+1}")

        # red rings for viol points
        for i in range(n_runs):
            run = runs[i]
            key = (run, f"Ctrl {lvl+1}")
            if key in viol:
                ax.scatter([x[i]], [np.clip(y[i], -3, 3)], s=120,
                           facecolors="none", edgecolors="red", linewidths=2.2, zorder=5)
                s = short_map.get(key, "")
                if s:
                    ax.text(x[i], np.clip(y[i], -3, 3)+0.15, s, color="red",
                            fontsize=8, ha="center", va="bottom")

    # Horizontal rules
    for y, lw, ls in [(0, 1.2, "-"),
                     (1, 0.9, "--"), (-1, 0.9, "--"),
                     (2, 0.9, "--"), (-2, 0.9, "--"),
                     (3, 1.2, "--"), (-3, 1.2, "--")]:
        ax.axhline(y, color="black", linewidth=lw, linestyle=ls, alpha=0.55)

    ax.set_title(title, fontsize=11)
    ax.set_ylabel("Z-score")
    ax.set_xlabel("Ngày / Lần")
    ax.set_xticks(x)
    # show fewer ticks if long
    if n_runs > 20:
        step = max(1, n_runs // 10)
        show = np.arange(0, n_runs, step)
        ax.set_xticks(show)
        ax.set_xticklabels([runs[i] for i in show], rotation=0, fontsize=8)
    else:
        ax.set_xticklabels(runs, rotation=0, fontsize=8)

    ax.set_ylim(-3.2, 3.2)
    ax.grid(True, alpha=0.25)
    ax.legend(loc="upper right", fontsize=8, frameon=True)
    fig.tight_layout()
    return fig


def _replace_placeholders_in_doc(doc: Document, mapping: dict):
    # paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                for run in p.runs:
                    if k in run.text:
                        run.text = run.text.replace(k, v)
    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for k, v in mapping.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)


def build_so_ghi_nhan_3muc_docx(export_df: pd.DataFrame,
                               z_df: pd.DataFrame,
                               point_df: Optional[pd.DataFrame],
                               meta: ReportMeta) -> io.BytesIO:
    """
    Tạo Word A4 cho 'Sổ ghi nhận & đánh giá 3 mức' + chèn biểu đồ L-J (ảnh).
    export_df: đã merge summary_df (có Trạng thái, Vi phạm loại bỏ, Người thực hiện)
    """
    # Create doc
    doc = Document()

    # Header/Footer (keep content placeholders similar to Excel form)
    apply_header_footer(
        doc,
        header_left=meta.don_vi,
        header_center="SỔ GHI NHẬN & ĐÁNH GIÁ KẾT QUẢ NỘI KIỂM",
        header_right="",
        version_text=meta.phien_ban,
        effective_date_text=meta.ngay_hieu_luc,
    )
    sec = doc.sections[0]
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2); sec.left_margin = Cm(2); sec.right_margin = Cm(2)

    title = doc.add_paragraph("SỔ GHI NHẬN & ĐÁNH GIÁ KẾT QUẢ NỘI KIỂM – 3 MỨC NỒNG ĐỘ")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.cell(0,0).text = "Tên xét nghiệm"; info.cell(0,1).text = meta.ten_xet_nghiem
    info.cell(1,0).text = "Thiết bị / Phương pháp"; info.cell(1,1).text = meta.thiet_bi_phuong_phap
    info.cell(2,0).text = "Lô QC / Hạn dùng"; info.cell(2,1).text = meta.lo_qc_han_dung
    info.cell(3,0).text = "Tháng / Năm"; info.cell(3,1).text = meta.thang_nam

    doc.add_paragraph("")
    p = doc.add_paragraph("BẢNG GHI NHẬN & ĐÁNH GIÁ THEO NGÀY")
    p.runs[0].bold = True

    # Ensure columns exist
    df = export_df.copy()
    for c in ["Ngày/Lần", "Ctrl 1", "Ctrl 2", "Ctrl 3",
              "z_Ctrl 1", "z_Ctrl 2", "z_Ctrl 3",
              "Trạng thái", "Vi phạm loại bỏ", "Người thực hiện"]:
        if c not in df.columns:
            df[c] = ""

    # build table
    tbl = doc.add_table(rows=1, cols=10)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Ngày", "L1", "L2", "L3", "Z-L1", "Z-L2", "Z-L3",
               "Đánh giá", "Quy tắc vi phạm", "Người thực hiện"]
    for i,h in enumerate(headers):
        tbl.cell(0,i).text = h

    for _, r in df.iterrows():
        cells = tbl.add_row().cells
        cells[0].text = _safe_str(r.get("Ngày/Lần",""))
        cells[1].text = _safe_str(r.get("Ctrl 1",""))
        cells[2].text = _safe_str(r.get("Ctrl 2",""))
        cells[3].text = _safe_str(r.get("Ctrl 3",""))
        cells[4].text = _safe_str(r.get("z_Ctrl 1",""))
        cells[5].text = _safe_str(r.get("z_Ctrl 2",""))
        cells[6].text = _safe_str(r.get("z_Ctrl 3",""))
        cells[7].text = _safe_str(r.get("Trạng thái",""))
        cells[8].text = _safe_str(r.get("Vi phạm loại bỏ",""))
        cells[9].text = _safe_str(r.get("Người thực hiện",""))

    doc.add_paragraph("")
    p = doc.add_paragraph("BIỂU ĐỒ LEVEY–JENNINGS (Z-SCORE)")
    p.runs[0].bold = True

    fig = build_lj_figure_from_z(z_df=z_df, point_df=point_df,
                                 title=f"{meta.ten_xet_nghiem} – Levey–Jennings (Z-score)")
    img_buf = io.BytesIO()
    fig.savefig(img_buf, format="png", dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    img_buf.seek(0)

    # Insert image ~ 3/4 A4 width (usable width ~ 17cm-4cm = 13cm)
    doc.add_picture(img_buf, width=Cm(12.5))

    doc.add_paragraph("")
    p = doc.add_paragraph("NHẬN XÉT – ĐÁNH GIÁ CHUNG")
    p.runs[0].bold = True
    doc.add_paragraph("{{NHAN_XET_CHUNG}}")

    doc.add_paragraph("")
    sign = doc.add_table(rows=1, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign.cell(0,0).text = "Người thực hiện\n(Ký, ghi rõ họ tên)"
    sign.cell(0,1).text = "Người kiểm tra / duyệt\n(Ký, ghi rõ họ tên)"

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out



def build_so_ghi_nhan_2muc_docx(export_df: pd.DataFrame,
                               z_df: pd.DataFrame,
                               point_df: Optional[pd.DataFrame],
                               meta: ReportMeta) -> io.BytesIO:
    """
    Tạo Word A4 cho 'Sổ ghi nhận & đánh giá 2 mức' + chèn biểu đồ L-J (ảnh).
    export_df: đã merge summary_df (có Trạng thái, Vi phạm loại bỏ, Người thực hiện)
    """
    doc = Document()

    apply_header_footer(
        doc,
        header_left=meta.don_vi,
        header_center="SỔ GHI NHẬN & ĐÁNH GIÁ KẾT QUẢ NỘI KIỂM",
        header_right="",
        version_text=meta.phien_ban,
        effective_date_text=meta.ngay_hieu_luc,
    )
    sec = doc.sections[0]
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2); sec.left_margin = Cm(2); sec.right_margin = Cm(2)

    title = doc.add_paragraph("SỔ GHI NHẬN & ĐÁNH GIÁ KẾT QUẢ NỘI KIỂM – 2 MỨC NỒNG ĐỘ")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(13)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # Info table
    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Tên xét nghiệm", meta.ten_xet_nghiem),
        ("Thiết bị / Phương pháp", meta.thiet_bi_phuong_phap),
        ("Lô QC / Hạn dùng", meta.lo_qc_han_dung),
        ("Tháng / Năm", meta.thang_nam),
    ]
    for i,(k,v) in enumerate(rows):
        info.cell(i,0).text = k
        info.cell(i,1).text = _safe_str(v)
    doc.add_paragraph("")

    df = export_df.copy()
    # headers
    tbl = doc.add_table(rows=1, cols=9)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Ngày", "L1", "L2", "Z-L1", "Z-L2",
               "Đánh giá", "Quy tắc vi phạm", "Người thực hiện", "Ghi chú"]
    for i,h in enumerate(headers):
        tbl.cell(0,i).text = h

    for _, r in df.iterrows():
        cells = tbl.add_row().cells
        cells[0].text = _safe_str(r.get("Ngày/Lần",""))
        cells[1].text = _safe_str(r.get("Ctrl 1",""))
        cells[2].text = _safe_str(r.get("Ctrl 2",""))
        cells[3].text = _safe_str(r.get("z_Ctrl 1",""))
        cells[4].text = _safe_str(r.get("z_Ctrl 2",""))
        cells[5].text = _safe_str(r.get("Trạng thái",""))
        cells[6].text = _safe_str(r.get("Vi phạm loại bỏ",""))
        cells[7].text = _safe_str(r.get("Người thực hiện",""))
        cells[8].text = _safe_str(r.get("Ghi chú",""))

    doc.add_paragraph("")
    doc.add_paragraph("BIỂU ĐỒ LEVEY–JENNINGS (Z-SCORE)").runs[0].bold = True

    fig = build_lj_figure_from_z(z_df=z_df, point_df=point_df, title="Levey–Jennings (Z-score)")
    img_buf = io.BytesIO()
    fig.savefig(img_buf, format="png", dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    img_buf.seek(0)
    # width ~ 3/4 A4 printable (approx 12.5cm)
    doc.add_picture(img_buf, width=Cm(12.5))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_cstk_3muc_docx(meta: ReportMeta, raw_df: Optional[pd.DataFrame], stats_df: pd.DataFrame) -> io.BytesIO:
    """
    Phiếu thiết lập CSTK – 3 mức.
    raw_df: DataFrame có cột ['L1','L2','L3'] (20–30 dòng). Có thể None.
    stats_df: DataFrame tổng hợp (theo page CSTK trong app), tối thiểu có cột:
        ['Control','Mean_X','SD_use','CV%_use'] hoặc tương đương.
    """
    doc = Document()

    apply_header_footer(
        doc,
        header_left=meta.don_vi,
        header_center="PHIẾU THIẾT LẬP CHỈ SỐ THỐNG KÊ (CSTK)",
        header_right="",
        version_text=meta.phien_ban,
        effective_date_text=meta.ngay_hieu_luc,
    )

    # Title
    title = doc.add_paragraph("PHIẾU THIẾT LẬP CHỈ SỐ THỐNG KÊ (CSTK) – 3 MỨC NỒNG ĐỘ")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(13)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # Info table
    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    pairs = [
        ("Tên xét nghiệm", meta.ten_xet_nghiem),
        ("Thiết bị / Phương pháp", meta.thiet_bi_phuong_phap),
        ("Lô QC / Hạn dùng", meta.lo_qc_han_dung),
        ("Ngày thiết lập", getattr(meta, "ngay_thiet_lap", "{{NGAY_THIET_LAP}}")),
    ]
    for i,(k,v) in enumerate(pairs):
        info.cell(i,0).text = k
        info.cell(i,1).text = v

    doc.add_paragraph("")

    # Raw data (gộp chung)
    doc.add_paragraph("BẢNG GIÁ TRỊ ĐO (RAW DATA)").runs[0].bold = True
    raw_tbl = doc.add_table(rows=1, cols=4)
    raw_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    raw_tbl.cell(0,0).text = "Lần đo"
    raw_tbl.cell(0,1).text = "Level 1"
    raw_tbl.cell(0,2).text = "Level 2"
    raw_tbl.cell(0,3).text = "Level 3"

    if raw_df is None or raw_df.empty:
        # placeholder row
        r = raw_tbl.add_row().cells
        r[0].text = "{{STT}}"
        r[1].text = "{{L1_VALUE}}"
        r[2].text = "{{L2_VALUE}}"
        r[3].text = "{{L3_VALUE}}"
    else:
        raw_df2 = raw_df.copy()
        # accept various column names
        colmap = {c.lower(): c for c in raw_df2.columns}
        def pick(*names):
            for n in names:
                if n in colmap:
                    return colmap[n]
            return None

        c1 = pick("l1","level1","level_1","lvl1")
        c2 = pick("l2","level2","level_2","lvl2")
        c3 = pick("l3","level3","level_3","lvl3")
        if c1 is None or c2 is None or c3 is None:
            # fallback: first 3 numeric cols
            num_cols = list(raw_df2.select_dtypes(include="number").columns)[:3]
            c1, c2, c3 = (num_cols + [None, None, None])[:3]

        for i_row, row in raw_df2.iterrows():
            rr = raw_tbl.add_row().cells
            rr[0].text = str(i_row + 1)
            rr[1].text = "" if c1 is None else ("" if pd.isna(row[c1]) else str(row[c1]))
            rr[2].text = "" if c2 is None else ("" if pd.isna(row[c2]) else str(row[c2]))
            rr[3].text = "" if c3 is None else ("" if pd.isna(row[c3]) else str(row[c3]))

    doc.add_paragraph("")

    # Thống kê theo mức: lấy từ stats_df nếu có các control L1/L2/L3
    doc.add_paragraph("TỔNG HỢP CHỈ SỐ THỐNG KÊ").runs[0].bold = True
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Mức QC", "Mean", "SD", "CV (%)", "Nguồn CSTK"]
    for i, h in enumerate(headers):
        tbl.cell(0, i).text = h

    # normalize stats_df
    s = stats_df.copy() if stats_df is not None else pd.DataFrame()
    if "Control" not in s.columns:
        # try to detect a control column
        for cand in ["control", "LEVEL", "level", "Muc", "Mức", "QC Level"]:
            if cand in s.columns:
                s = s.rename(columns={cand: "Control"})
                break

    def get_val(ctrl_key: str, key_candidates):
        if s is None or s.empty or "Control" not in s.columns:
            return ""
        row = s[s["Control"].astype(str).str.lower().isin([ctrl_key.lower(), f"ctrl {ctrl_key[-1]}".lower(), f"level {ctrl_key[-1]}".lower()])]
        if row.empty:
            # try exact
            row = s[s["Control"].astype(str).str.strip().str.lower() == ctrl_key.lower()]
        if row.empty:
            return ""
        for k in key_candidates:
            if k in s.columns:
                v = row.iloc[0][k]
                return "" if pd.isna(v) else str(v)
        return ""

    for ctrl, tag in [("L1", "Level 1"), ("L2", "Level 2"), ("L3", "Level 3")]:
        r = tbl.add_row().cells
        r[0].text = tag
        r[1].text = get_val(ctrl, ["Mean_X", "Mean", "mean"])
        r[2].text = get_val(ctrl, ["SD_use", "SD", "sd"])
        r[3].text = get_val(ctrl, ["CV%_use", "CV%", "CV", "cv"])
        r[4].text = getattr(meta, "nguon_cstk", "{{NGUON_CSTK}}")

    doc.add_paragraph("")
    doc.add_paragraph("Nhận xét: {{NHAN_XET}}")

    doc.add_paragraph("")
    sign = doc.add_table(rows=1, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign.cell(0,0).text = "Người lập\n(Ký, ghi rõ họ tên)"
    sign.cell(0,1).text = "Người duyệt\n(Ký, ghi rõ họ tên)"

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
        
def build_cstk_2muc_docx(meta: ReportMeta, raw_df: Optional[pd.DataFrame], stats_df: pd.DataFrame) -> io.BytesIO:
    """
    Phiếu thiết lập CSTK – 2 mức.
    raw_df: DataFrame có cột ['L1','L2'] (20–30 dòng). Có thể None.
    stats_df: DataFrame tổng hợp, tối thiểu có cột:
        ['Control','Mean_X','SD_use','CV%_use'] hoặc tương đương.
    """
    doc = Document()

    apply_header_footer(
        doc,
        header_left=meta.don_vi,
        header_center="PHIẾU THIẾT LẬP CHỈ SỐ THỐNG KÊ (CSTK)",
        header_right="",
        version_text=meta.phien_ban,
        effective_date_text=meta.ngay_hieu_luc,
    )

    title = doc.add_paragraph("PHIẾU THIẾT LẬP CHỈ SỐ THỐNG KÊ (CSTK) – 2 MỨC NỒNG ĐỘ")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(13)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    info = doc.add_table(rows=3, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Tên xét nghiệm", meta.ten_xet_nghiem),
        ("Thiết bị / Phương pháp", meta.thiet_bi_phuong_phap),
        ("Lô QC / Hạn dùng", meta.lo_qc_han_dung),
    ]
    for i,(k,v) in enumerate(rows):
        info.cell(i,0).text = k
        info.cell(i,1).text = _safe_str(v)
    doc.add_paragraph("")

    # Raw data (optional)
    if raw_df is not None and not raw_df.empty:
        doc.add_paragraph("DỮ LIỆU GỐC (RAW DATA)").runs[0].bold = True
        t = doc.add_table(rows=1, cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.cell(0,0).text = "STT"
        t.cell(0,1).text = "L1"
        t.cell(0,2).text = "L2"
        for i in range(len(raw_df)):
            rr = t.add_row().cells
            rr[0].text = str(i+1)
            rr[1].text = _safe_str(raw_df.iloc[i].get("L1",""))
            rr[2].text = _safe_str(raw_df.iloc[i].get("L2",""))
        doc.add_paragraph("")

    # Stats summary
    doc.add_paragraph("KẾT QUẢ TÍNH TOÁN CSTK").runs[0].bold = True

    s = stats_df.copy()
    # try normalize expected column names
    col_mean = "Mean_X" if "Mean_X" in s.columns else ("Mean" if "Mean" in s.columns else None)
    col_sd = "SD_use" if "SD_use" in s.columns else ("SD" if "SD" in s.columns else None)
    col_cv = "CV%_use" if "CV%_use" in s.columns else ("CV%" if "CV%" in s.columns else None)

    t2 = doc.add_table(rows=1, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.cell(0,0).text = "Mức QC"
    t2.cell(0,1).text = "Mean"
    t2.cell(0,2).text = "SD"
    t2.cell(0,3).text = "CV (%)"

    # Only Ctrl 1 & 2
    for ctrl_label in ["Ctrl 1", "Ctrl 2"]:
        row = s[s["Control"] == ctrl_label]
        mean_v = row.iloc[0][col_mean] if (not row.empty and col_mean) else ""
        sd_v = row.iloc[0][col_sd] if (not row.empty and col_sd) else ""
        cv_v = row.iloc[0][col_cv] if (not row.empty and col_cv) else ""
        rr = t2.add_row().cells
        rr[0].text = ctrl_label.replace("Ctrl ", "Level ")
        rr[1].text = _safe_str(mean_v)
        rr[2].text = _safe_str(sd_v)
        rr[3].text = _safe_str(cv_v)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


